"""Regression tests for auth hardening, text metrics and bounded batches."""
from __future__ import annotations

import uuid
from io import BytesIO
from datetime import datetime, timedelta, timezone

import pytest
from fastapi.testclient import TestClient

from app.core.rate_limit import limiter
from app.main import app
from app.db import connect, encode, new_id
from app.core.embeddings import hash_embedding
from app.services.knowledge_hub import (
    EMBEDDING_DIMENSION,
    _chunk_text,
    rebuild_item_embeddings,
    search,
)
from app.services.knowledge_parser import extract_document_text
from app.services.text_metrics import count_content_chars


@pytest.fixture(autouse=True)
def _reset_rate_limits():
    limiter.reset()


def test_content_character_count_ignores_whitespace():
    assert count_content_chars("第一章\n\n 你好 world\t!") == 11


def test_chunk_text_is_stable_and_overlapping():
    text = "x" * 1700
    chunks = _chunk_text(text)
    assert [len(chunk) for chunk in chunks] == [800, 800, 300]
    assert chunks == _chunk_text(text)


def test_local_embedding_is_deterministic_and_normalized():
    vector = hash_embedding("这是需要建立索引的世界观文本")
    assert len(vector) == EMBEDDING_DIMENSION
    assert vector == hash_embedding("这是需要建立索引的世界观文本")
    assert abs(sum(value * value for value in vector) - 1.0) < 1e-6


def test_refresh_cookie_requires_csrf_and_rotates():
    # Production deliberately marks the refresh cookie Secure.  Use an HTTPS
    # test origin so TestClient exercises the real cookie/CSRF path instead of
    # silently dropping the cookie on an http:// request.
    with TestClient(app, base_url="https://testserver") as client:
        email = f"refresh-{uuid.uuid4().hex}@nc.dev"
        registered = client.post("/api/v1/auth/register", json={"email": email, "password": "test1234"})
        assert registered.status_code == 200
        old_refresh = client.cookies.get("refresh_token")
        csrf = client.cookies.get("csrf_token")
        assert old_refresh and csrf
        assert "HttpOnly" in registered.headers.get("set-cookie", "")
        assert "refresh_token" not in registered.json()["data"]

        rejected = client.post("/api/v1/auth/refresh")
        assert rejected.status_code == 403

        refreshed = client.post("/api/v1/auth/refresh", headers={"X-CSRF-Token": csrf})
        assert refreshed.status_code == 200
        assert refreshed.json()["data"]["access_token"]
        assert client.cookies.get("refresh_token") != old_refresh

    with TestClient(app, base_url="https://testserver") as replay_client:
        replay = replay_client.post("/api/v1/auth/refresh", json={"refresh_token": old_refresh})
        assert replay.status_code == 401


def test_batch_chapter_count_is_bounded():
    with TestClient(app) as client:
        email = f"batch-{uuid.uuid4().hex}@nc.dev"
        registered = client.post("/api/v1/auth/register", json={"email": email, "password": "test1234"})
        token = registered.json()["data"]["access_token"]
        csrf = client.cookies.get("csrf_token")
        response = client.post(
            f"/api/v1/novels/{uuid.uuid4()}/chapters/batch",
            json={"chapter_count": 51},
            headers={"Authorization": f"Bearer {token}", "X-CSRF-Token": csrf},
        )
        assert response.status_code == 422


def test_knowledge_reindex_replaces_old_chunks(monkeypatch):
    monkeypatch.setenv("EMBEDDING_BACKEND", "hash")
    monkeypatch.setenv("NOVELCRAFT_ALLOW_HASH_EMBEDDING", "true")
    db = connect()
    project = db.execute("SELECT id FROM projects LIMIT 1").fetchone()
    item_id = new_id()
    unique_term = f"量子航道-{uuid.uuid4().hex}"
    db.execute(
        "INSERT INTO knowledge_items (id, project_id, kind, title, body) VALUES (%s, %s, 'worldview', %s, %s)",
        (item_id, project["id"], "星港世界观", f"星港采用{unique_term}。" * 120),
    )
    db.commit()
    db.close()

    first_count = rebuild_item_embeddings(item_id)
    second_count = rebuild_item_embeddings(item_id)
    db = connect()
    stored_count = db.execute("SELECT COUNT(*) AS total FROM knowledge_vectors WHERE item_id = %s", (item_id,)).fetchone()["total"]
    db.close()
    assert first_count == second_count == stored_count
    assert str(search(unique_term, project_id=str(project["id"]))[0]["id"]) == item_id


def test_batch_can_be_created_and_cancelled(monkeypatch):
    class FakeTask:
        id = "fake-celery-task"

    from app.workers.tasks import batch_generate_chapters_task
    monkeypatch.setattr(batch_generate_chapters_task, "delay", lambda *args, **kwargs: FakeTask())

    with TestClient(app) as client:
        email = f"batch-flow-{uuid.uuid4().hex}@nc.dev"
        registered = client.post("/api/v1/auth/register", json={"email": email, "password": "test1234"})
        token = registered.json()["data"]["access_token"]
        headers = {"Authorization": f"Bearer {token}"}
        project_id = client.get("/api/v1/projects", headers=headers).json()["data"][0]["id"]
        novel = client.post(
            f"/api/v1/projects/{project_id}/novels",
            headers=headers,
            json={"idea": "一座城市每天遗忘一个人", "target_words": 100000},
        ).json()["data"]
        created = client.post(
            f"/api/v1/novels/{novel['id']}/chapters/batch",
            headers=headers,
            json={"chapter_count": 3},
        )
        assert created.status_code == 200
        batch_id = created.json()["data"]["batch_id"]
        cancelled = client.post(f"/api/v1/generation-batches/{batch_id}/cancel", headers=headers)
        assert cancelled.status_code == 200
        state = client.get(f"/api/v1/generation-batches/{batch_id}", headers=headers).json()["data"]
        assert state["status"] == "cancelled"
        assert state["cancel_requested"] is True


def test_chapter_generation_releases_lock_on_failure(monkeypatch):
    from app.workers import lock
    from app.workers import tasks

    released: list[str] = []
    monkeypatch.setattr(lock, "acquire_lock", lambda key: True)
    monkeypatch.setattr(lock, "release_lock", lambda key: released.append(key))
    monkeypatch.setattr(tasks, "_generate_next_chapter_unlocked", lambda *_args: (_ for _ in ()).throw(RuntimeError("boom")))

    try:
        tasks.gen_next_chapter_task.run("novel-id", "project-id")
    except RuntimeError as exc:
        assert str(exc) == "boom"
    else:
        raise AssertionError("generation failure was not propagated")
    assert released == ["lock:novel:novel-id:gen_chapter"]


def test_offline_content_sync_is_idempotent_and_preserves_conflicts():
    with TestClient(app) as client:
        email = f"offline-sync-{uuid.uuid4().hex}@nc.dev"
        registered = client.post("/api/v1/auth/register", json={"email": email, "password": "test1234"})
        headers = {"Authorization": f"Bearer {registered.json()['data']['access_token']}"}
        project_id = client.get("/api/v1/projects", headers=headers).json()["data"][0]["id"]
        novel = client.post(
            f"/api/v1/projects/{project_id}/novels",
            headers=headers,
            json={"idea": "离线冲突测试小说", "target_words": 100000},
        ).json()["data"]
        original_updated_at = novel["updated_at"]
        first_mutation = f"mutation-{uuid.uuid4()}"
        first_body = {"type": "doc", "content": [{"type": "paragraph", "text": "联网版本"}]}
        applied = client.put(
            f"/api/v1/contents/{novel['id']}",
            headers=headers,
            json={
                "body": first_body,
                "base_updated_at": original_updated_at,
                "client_mutation_id": first_mutation,
            },
        ).json()["data"]
        assert applied["sync_status"] == "applied"

        replayed = client.put(
            f"/api/v1/contents/{novel['id']}",
            headers=headers,
            json={
                "body": first_body,
                "base_updated_at": original_updated_at,
                "client_mutation_id": first_mutation,
            },
        ).json()["data"]
        assert replayed["mutation_replayed"] is True

        conflict = client.put(
            f"/api/v1/contents/{novel['id']}",
            headers=headers,
            json={
                "body": {"type": "doc", "content": [{"type": "paragraph", "text": "离线版本"}]},
                "base_updated_at": original_updated_at,
                "client_mutation_id": f"mutation-{uuid.uuid4()}",
            },
        ).json()["data"]
        assert conflict["sync_status"] == "conflict"
        assert conflict["conflict_version_id"]
        assert conflict["body"] == first_body
        resolved = client.put(
            f"/api/v1/contents/{novel['id']}",
            headers=headers,
            json={
                "body": {"type": "doc", "content": [{"type": "paragraph", "text": "合并版本"}]},
                "base_updated_at": conflict["updated_at"],
                "client_mutation_id": f"mutation-{uuid.uuid4()}",
            },
        ).json()["data"]
        assert resolved["sync_status"] == "applied"
        db = connect()
        conflict_version = db.execute(
            "SELECT reason FROM versions WHERE id = %s",
            (conflict["conflict_version_id"],),
        ).fetchone()
        db.close()
        assert conflict_version["reason"] == "offline_conflict_resolved"


def test_offline_ai_mutation_replay_returns_cached_result():
    from app.gateway import complete

    mutation_id = f"ai-mutation-{uuid.uuid4()}"
    expected = {"text": "已经生成且不会重复计费"}
    db = connect()
    project = db.execute("SELECT id FROM projects LIMIT 1").fetchone()
    db.execute(
        """
        INSERT INTO ai_calls (
            id, provider, model, prompt_name, task_type, input, output,
            status, client_mutation_id, project_id
        ) VALUES (%s, 'test', 'test', 'editor.polish', 'editor_polish', %s, %s, 'succeeded', %s, %s)
        """,
        (new_id(), encode({}), encode(expected), mutation_id, project["id"]),
    )
    db.commit()
    db.close()
    result = complete(
        run_id=None,
        node_key=None,
        project_id=str(project["id"]),
        task_type="editor_polish",
        prompt_name="editor.polish",
        variables={"selection": "原文"},
        client_mutation_id=mutation_id,
    )
    assert result == expected


def test_docx_extraction_and_project_scoped_knowledge_import(monkeypatch):
    monkeypatch.setenv("EMBEDDING_BACKEND", "hash")
    monkeypatch.setenv("NOVELCRAFT_ALLOW_HASH_EMBEDDING", "true")
    from docx import Document

    document = Document()
    document.add_heading("世界规则", level=1)
    document.add_paragraph("量子航道只能在月食期间开启。")
    buffer = BytesIO()
    document.save(buffer)
    assert "量子航道" in extract_document_text(buffer.getvalue(), "world.docx")

    with TestClient(app) as client:
        owner_email = f"knowledge-owner-{uuid.uuid4().hex}@nc.dev"
        owner = client.post("/api/v1/auth/register", json={"email": owner_email, "password": "test1234"}).json()["data"]
        owner_headers = {"Authorization": f"Bearer {owner['access_token']}"}
        project_id = client.get("/api/v1/projects", headers=owner_headers).json()["data"][0]["id"]
        imported = client.post(
            f"/api/v1/knowledge/import?project_id={project_id}",
            headers=owner_headers,
            files={"file": ("world.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert imported.status_code == 200
        assert imported.json()["data"]["imported"] >= 1
        assert imported.json()["data"]["chunks"] >= 1

        other_email = f"knowledge-other-{uuid.uuid4().hex}@nc.dev"
        other = client.post("/api/v1/auth/register", json={"email": other_email, "password": "test1234"}).json()["data"]
        rejected = client.post(
            f"/api/v1/knowledge/import?project_id={project_id}",
            headers={"Authorization": f"Bearer {other['access_token']}"},
            files={"file": ("note.md", b"# blocked", "text/markdown")},
        )
        assert rejected.status_code == 403


def test_login_failures_lock_account_until_counter_is_cleared():
    from app.core.security import clear_login_failures

    with TestClient(app) as client:
        email = f"locked-{uuid.uuid4().hex}@nc.dev"
        client.post("/api/v1/auth/register", json={"email": email, "password": "correct-password"})
        for _ in range(5):
            failed = client.post("/api/v1/auth/login", json={"email": email, "password": "wrong-password"})
            assert failed.status_code == 401
        locked = client.post("/api/v1/auth/login", json={"email": email, "password": "correct-password"})
        assert locked.status_code == 429
        clear_login_failures(email)
        accepted = client.post("/api/v1/auth/login", json={"email": email, "password": "correct-password"})
        assert accepted.status_code == 200


def test_publish_schedule_and_analytics_are_project_scoped():
    with TestClient(app) as client:
        owner_email = f"publish-owner-{uuid.uuid4().hex}@nc.dev"
        owner = client.post("/api/v1/auth/register", json={"email": owner_email, "password": "test1234"}).json()["data"]
        owner_headers = {"Authorization": f"Bearer {owner['access_token']}"}
        project_id = client.get("/api/v1/projects", headers=owner_headers).json()["data"][0]["id"]
        novel = client.post(
            f"/api/v1/projects/{project_id}/novels",
            headers=owner_headers,
            json={"idea": "发布权限测试小说", "target_words": 100000},
        ).json()["data"]
        scheduled_at = (datetime.now(timezone.utc) + timedelta(days=1)).isoformat()
        scheduled = client.post(
            "/api/v1/publish/schedule",
            headers=owner_headers,
            params={"content_id": novel["id"], "platform": "medium", "scheduled_at": scheduled_at},
        )
        assert scheduled.status_code == 200
        assert len(client.get("/api/v1/publish/schedule", headers=owner_headers).json()["data"]) >= 1
        assert client.get("/api/v1/analytics/roi", headers=owner_headers, params={"project_id": project_id}).status_code == 200

        other_email = f"publish-other-{uuid.uuid4().hex}@nc.dev"
        other = client.post("/api/v1/auth/register", json={"email": other_email, "password": "test1234"}).json()["data"]
        other_headers = {"Authorization": f"Bearer {other['access_token']}"}
        forbidden = client.post(
            "/api/v1/publish/schedule",
            headers=other_headers,
            params={"content_id": novel["id"], "platform": "medium", "scheduled_at": scheduled_at},
        )
        assert forbidden.status_code == 403
        assert client.get("/api/v1/analytics/roi", headers=other_headers, params={"project_id": project_id}).status_code == 403


def test_publish_modes_and_wordpress_target_are_restricted():
    from app.workers.m4_tasks import _is_public_https_url

    assert _is_public_https_url("http://example.com") is False
    assert _is_public_https_url("https://127.0.0.1") is False
    with TestClient(app) as client:
        email = f"publish-mode-{uuid.uuid4().hex}@nc.dev"
        registered = client.post("/api/v1/auth/register", json={"email": email, "password": "test1234"}).json()["data"]
        headers = {"Authorization": f"Bearer {registered['access_token']}"}
        project_id = client.get("/api/v1/projects", headers=headers).json()["data"][0]["id"]
        novel = client.post(
            f"/api/v1/projects/{project_id}/novels",
            headers=headers,
            json={"idea": "发布模式测试小说", "target_words": 100000},
        ).json()["data"]
        rejected = client.post(
            "/api/v1/publish",
            headers=headers,
            params={"content_id": novel["id"], "platform": "kdp", "mode": "auto"},
        )
        assert rejected.status_code == 400
